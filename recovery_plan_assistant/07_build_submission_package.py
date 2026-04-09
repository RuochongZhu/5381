from __future__ import annotations

import os
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches

from shared import BASE_DIR, OUTPUT_DIR, SCREENSHOT_DIR, render_terminal_screenshot

SCRIPT_SEQUENCE = [
    ("01_basic_function_calling.py", "basic_function_calling"),
    ("02_agent_wrapper_tool_demo.py", "agent_wrapper_tool_demo"),
    ("03_multi_agent_with_tools.py", "multi_agent_with_tools"),
    ("04_prompt_design_multi_agent.py", "prompt_design_multi_agent"),
    ("05_custom_rag_query.py", "custom_rag_query"),
    ("06_final_ai_agent_system.py", "final_ai_agent_system"),
]


def run_script(script_name: str) -> str:
    env = os.environ.copy()
    env.setdefault("DSAI_MODEL", "qwen2.5:3b")
    result = subprocess.run(
        [sys.executable, script_name],
        cwd=BASE_DIR,
        capture_output=True,
        text=True,
        env=env,
        check=True,
    )
    return result.stdout.strip()


def build_docx() -> Path:
    repo_url = os.environ.get("SUBMISSION_REPO_URL", "").rstrip("/")
    doc = Document()
    doc.add_heading("Homework 2: Recovery Plan Assistant", level=0)

    p = doc.add_paragraph()
    p.add_run("Important: ").bold = True
    p.add_run(
        "The writing component below is intentionally left as a fill-in template because the assignment says it must be written in your own words."
    )

    doc.add_heading("Writing Component", level=1)
    doc.add_paragraph("[Replace this paragraph with your own explanation of what the system does.]")
    doc.add_paragraph("[Replace this paragraph with your own explanation of how the agents, RAG, and tools work together.]")
    doc.add_paragraph("[Replace this paragraph with your own explanation of design choices, limitations, or challenges.]")

    doc.add_heading("Code Links", level=1)
    if repo_url:
        doc.add_paragraph(f"Multi-agent orchestration: {repo_url}/blob/main/student_submissions/recovery_plan_assistant/04_prompt_design_multi_agent.py")
        doc.add_paragraph(f"RAG implementation: {repo_url}/blob/main/student_submissions/recovery_plan_assistant/05_custom_rag_query.py")
        doc.add_paragraph(f"Function calling and tools: {repo_url}/blob/main/student_submissions/recovery_plan_assistant/shared.py")
        doc.add_paragraph(f"Main system file: {repo_url}/blob/main/student_submissions/recovery_plan_assistant/06_final_ai_agent_system.py")
    else:
        doc.add_paragraph("Set SUBMISSION_REPO_URL before running this script to auto-fill GitHub links.")
        doc.add_paragraph("Example: export SUBMISSION_REPO_URL=https://github.com/YOUR_USERNAME/YOUR_REPO")
        doc.add_paragraph("Files to link:")
        doc.add_paragraph("student_submissions/recovery_plan_assistant/04_prompt_design_multi_agent.py")
        doc.add_paragraph("student_submissions/recovery_plan_assistant/05_custom_rag_query.py")
        doc.add_paragraph("student_submissions/recovery_plan_assistant/shared.py")
        doc.add_paragraph("student_submissions/recovery_plan_assistant/06_final_ai_agent_system.py")

    doc.add_heading("Screenshots / Outputs", level=1)
    for name in [
        "multi_agent_with_tools.png",
        "custom_rag_query.png",
        "final_ai_agent_system.png",
        "mcp_tool_call.png",
    ]:
        image_path = SCREENSHOT_DIR / name
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(6.5))

    doc.add_heading("Documentation", level=1)
    doc.add_paragraph(
        "System architecture: Agent 1 retrieves evidence from a local recovery dataset, Agent 2 scores candidates, and Agent 3 writes the final brief."
    )
    doc.add_paragraph(
        "RAG data source: a curated CSV of twelve recovery projects with fields for community, sector, issue, action, and priority. The search function ranks rows by keyword overlap and filters by community or sector when provided."
    )
    doc.add_paragraph(
        "Tool functions: search_recovery_projects(query, community, sector, top_n) returns matching project records as JSON; summarize_sector_counts(community) returns sector counts as JSON; calculate_priority_score(impact, feasibility, urgency) returns a weighted score and label."
    )
    doc.add_paragraph(
        "Technical details: Python 3.9 virtual environment, local Ollama server on http://127.0.0.1:11434, model qwen2.5:3b, and optional FastAPI + uvicorn for the MCP server."
    )
    doc.add_paragraph(
        "Usage instructions: activate .venv, start Ollama, run the scripts inside student_submissions/recovery_plan_assistant, then use 07_build_submission_package.py to regenerate output logs, screenshots, and the .docx file."
    )

    doc_path = BASE_DIR / "Homework2_Recovery_Plan_Assistant.docx"
    doc.save(doc_path)
    return doc_path


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)

    for script_name, screenshot_stem in SCRIPT_SEQUENCE:
        stdout = run_script(script_name)
        render_terminal_screenshot(
            title=script_name,
            body_text=stdout,
            output_path=SCREENSHOT_DIR / f"{screenshot_stem}.png",
        )

    mcp_output = OUTPUT_DIR / "07_mcp_test.txt"
    if mcp_output.exists():
        render_terminal_screenshot(
            title="mcp/test_mcp.py",
            body_text=mcp_output.read_text(encoding="utf-8"),
            output_path=SCREENSHOT_DIR / "mcp_tool_call.png",
        )

    doc_path = build_docx()
    print(f"Built screenshots in {SCREENSHOT_DIR}")
    print(f"Built docx at {doc_path}")


if __name__ == "__main__":
    main()
